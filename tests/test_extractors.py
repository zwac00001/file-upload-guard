import os
import subprocess
import sys
from pathlib import Path

from docx import Document

from file_upload_guard.extractors import extract_text_from_path


def test_extract_text_from_docx_reads_paragraphs(tmp_path: Path):
    path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Quarterly report")
    document.add_paragraph("Revenue increased")
    document.save(path)

    result = extract_text_from_path(path)

    assert result.success is True
    assert "Quarterly report" in result.text
    assert "Revenue increased" in result.text


def test_extract_text_from_unknown_suffix_reports_failure(tmp_path: Path):
    path = tmp_path / "archive.zip"
    path.write_bytes(b"PK\x03\x04")

    result = extract_text_from_path(path)

    assert result.success is False
    assert result.reason == "No extractor available for .zip"


def test_cli_import_succeeds_without_optional_pdf_or_docx_dependencies(
    tmp_path: Path,
):
    blocker = tmp_path / "sitecustomize.py"
    blocker.write_text(
        "\n".join(
            [
                "import builtins",
                "_real_import = builtins.__import__",
                "",
                "def _blocked_import(name, globals=None, locals=None, fromlist=(), level=0):",
                "    if name in {'docx', 'pypdf'}:",
                "        raise ImportError(f'blocked optional dependency: {name}')",
                "    return _real_import(name, globals, locals, fromlist, level)",
                "",
                "builtins.__import__ = _blocked_import",
            ]
        ),
        encoding="utf-8",
    )

    env = os.environ.copy()
    env["PYTHONPATH"] = os.pathsep.join(
        [str(tmp_path), str(Path("src").resolve()), env.get("PYTHONPATH", "")]
    ).rstrip(os.pathsep)

    result = subprocess.run(
        [sys.executable, "-c", "import file_upload_guard.cli; print('ok')"],
        cwd=Path.cwd(),
        env=env,
        capture_output=True,
        text=True,
    )

    assert result.returncode == 0, result.stderr
    assert result.stdout.strip() == "ok"
